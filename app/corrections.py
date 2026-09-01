"""Import client-corrected Excel/Word compliance sheets into gold + rule book."""

from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from openpyxl import load_workbook

from app.compliance import (
    CATALOG_NAME,
    load_patterns,
    load_schneider_sources,
    match_capability,
    resolve_pattern_key,
    shorten_remark,
)
from app.gold_reference import append_reference_rows

ROOT = Path(__file__).resolve().parent.parent
GOLD_DIR = ROOT / "data" / "gold"
GOLD_LOG = GOLD_DIR / "corrections.jsonl"
REFERENCE_DB = GOLD_DIR / "reference_answers.jsonl"
PATTERNS_PATH = ROOT / "data" / "patterns" / "drive_capabilities.json"
SOURCES_PATH = ROOT / "data" / "patterns" / "schneider_sources.json"

STATUS_MAP = {
    "YES": "yes",
    "Y": "yes",
    "COMPLY": "yes",
    "COMPLIES": "yes",
    "NO": "no",
    "N": "no",
    "DOES NOT COMPLY": "no",
    "N/A": "na",
    "NA": "na",
    "NOT APPLICABLE": "na",
}


def _norm_status(raw: str) -> str | None:
    t = re.sub(r"\s+", " ", (raw or "")).strip().upper()
    t = t.replace("COMPLY (YES)", "YES").replace("DOES NOT COMPLY (NO)", "NO")
    t = t.replace("NOT APPLICABLE (N/A)", "N/A")
    if t.startswith("YES") or t.startswith("COMPLY"):
        return "yes"
    if t.startswith("NO") or "DOES NOT" in t:
        return "no"
    if "N/A" in t or t in ("NA", "NOT APPLICABLE"):
        return "na"
    return STATUS_MAP.get(t)


def _parse_xlsx(data: bytes) -> list[dict[str, str]]:
    wb = load_workbook(BytesIO(data), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    header = [str(c or "").strip().lower() for c in rows[0]]
    # Find our export header; skip title rows
    start = 0
    for i, row in enumerate(rows[:12]):
        joined = " ".join(str(c or "").lower() for c in row)
        if "requirement" in joined and ("complies" in joined or "why" in joined):
            header = [str(c or "").strip().lower() for c in row]
            start = i + 1
            break
        # Client gold sheets: Clause / Comment / Reason
        if "clause" in joined and ("comment" in joined or "comply" in joined):
            header = [str(c or "").strip().lower() for c in row]
            start = i + 1
            break

    def col(*names: str) -> int | None:
        for n in names:
            for idx, h in enumerate(header):
                if n in h:
                    return idx
        return None

    i_req = col("requirement", "clause", "specs")
    i_status = col("complies", "comply", "yes/no", "comment")
    i_why = col("why", "remark", "reason")
    i_src = col("schneider source", "source")
    i_clause = col("clause")
    # If Comment holds Comply and Reason holds remark, treat Comment as status
    out = []
    for row in rows[start:]:
        if not row:
            continue
        req = str(row[i_req] if i_req is not None and i_req < len(row) else "").strip()
        if not req or req.lower() in ("requirement text", "requirement", "clause", "clause "):
            continue
        status_raw = str(row[i_status] if i_status is not None and i_status < len(row) else "").strip()
        why = str(row[i_why] if i_why is not None and i_why < len(row) else "").strip()
        # Comment column often is just "Comply" — if status empty but why looks like Comply…
        status = status_raw
        if not _norm_status(status) and _norm_status(why):
            status, why = why, status_raw
        # If Comment is "Comply" and Reason has the long remark — good
        if _norm_status(status_raw) and why:
            status = status_raw
        src = str(row[i_src] if i_src is not None and i_src < len(row) else "").strip()
        clause = str(row[i_clause] if i_clause is not None and i_clause < len(row) else "").strip()
        if not _norm_status(status) and not why:
            continue
        out.append(
            {
                "clause_id": clause if clause != req else "",
                "requirement": req,
                "status": status,
                "remarks": why if why != status else "",
                "source_document": src,
            }
        )
    return out


def _parse_docx_tables(doc) -> list[dict[str, str]]:
    """Gold Word sheets often store Specs | Reply in tables."""
    items: list[dict[str, str]] = []
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        joined = " ".join(header)
        # Specs Points | Schneider Electric Comment
        # Item | Specs | Electra Control Reply
        i_req = None
        i_reply = None
        for idx, h in enumerate(header):
            if any(k in h for k in ("specs", "requirement", "clause", "point")):
                i_req = idx
            if any(k in h for k in ("reply", "comment", "schneider", "comply")):
                i_reply = idx
        if i_req is None or i_reply is None:
            # two-column specs/comment without clear header
            if len(header) >= 2 and (
                "specs" in joined or "comment" in joined or "point" in joined
            ):
                i_req, i_reply = 0, 1
            else:
                # Try all rows as requirement | reply when 2–3 cols
                ncols = len(table.rows[0].cells)
                if ncols < 2:
                    continue
                i_req = 0 if ncols == 2 else 1
                i_reply = ncols - 1
                start = 0
                for ri, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) <= i_reply:
                        continue
                    req, reply = cells[i_req], cells[i_reply]
                    if not req or len(req) < 12:
                        continue
                    if not reply:
                        continue
                    status = _norm_status(reply) or (
                        "yes" if re.match(r"(?i)^\s*comply", reply) else None
                    )
                    if not status and len(reply) > 8:
                        status = "yes"
                        remark = reply
                    else:
                        remark = re.sub(r"(?i)^\s*comply\s*[.—:\-]*\s*", "", reply).strip()
                    if not status:
                        continue
                    items.append(
                        {
                            "clause_id": "",
                            "requirement": req[:500],
                            "status": status,
                            "remarks": remark[:400],
                            "source_document": "",
                        }
                    )
                continue

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(i_req, i_reply):
                continue
            req = cells[i_req]
            reply = cells[i_reply]
            if not req or len(req) < 12:
                continue
            if not reply:
                continue
            status = _norm_status(reply)
            if not status and re.match(r"(?i)^\s*comply", reply):
                status = "yes"
            if not status and len(reply) > 15:
                status = "yes"
            if not status:
                continue
            remark = re.sub(r"(?i)^\s*comply\s*[.—:\-]*\s*", "", reply).strip()
            items.append(
                {
                    "clause_id": "",
                    "requirement": req[:500],
                    "status": status,
                    "remarks": remark[:400],
                    "source_document": "",
                }
            )
    return items


def _parse_docx(data: bytes) -> list[dict[str, str]]:
    from docx import Document

    doc = Document(BytesIO(data))
    table_items = _parse_docx_tables(doc)
    if table_items:
        return table_items

    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    items: list[dict[str, str]] = []
    i = 0
    while i < len(paras):
        line = paras[i]
        nxt = paras[i + 1] if i + 1 < len(paras) else ""
        if re.match(
            r"^(Comply \(Yes\)|Does not comply \(No\)|Not applicable \(N/A\))",
            nxt,
            re.I,
        ):
            status_line = nxt
            remarks = ""
            m = re.match(
                r"^(Comply \(Yes\)|Does not comply \(No\)|Not applicable \(N/A\))\s*[—\-–:]?\s*(.*)$",
                status_line,
                re.I,
            )
            if m:
                status_raw, remarks = m.group(1), m.group(2).strip()
            else:
                status_raw = status_line
            source = ""
            j = i + 2
            while j < len(paras) and j <= i + 5:
                if paras[j].lower().startswith("schneider source"):
                    source = re.sub(r"(?i)^schneider source:\s*", "", paras[j]).strip()
                    break
                if paras[j].lower().startswith("consultant page") and "schneider source" in paras[j].lower():
                    mm = re.search(r"(?i)schneider source:\s*(.+)$", paras[j])
                    source = mm.group(1).strip() if mm else ""
                    break
                j += 1
            clause = ""
            req = line
            cm = re.match(r"^(\S+)\s{2,}(.+)$", line)
            if cm:
                clause, req = cm.group(1), cm.group(2)
            items.append(
                {
                    "clause_id": clause,
                    "requirement": req,
                    "status": status_raw,
                    "remarks": remarks,
                    "source_document": source,
                }
            )
            i = j + 1
            continue
        i += 1
    return items


def parse_correction_file(filename: str, data: bytes) -> list[dict[str, str]]:
    name = (filename or "").lower()
    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        return _parse_xlsx(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    raise ValueError("Please upload an Excel (.xlsx) or Word (.docx) file.")


def _extract_page(source: str) -> str | None:
    m = re.search(r"(?i)\bp\.?\s*([0-9]+(?:/[0-9]+)?)\b", source or "")
    return m.group(1) if m else None


def _doc_kind(source: str) -> str | None:
    """Only Catalog / Programming / Installation count as Schneider source docs."""
    t = (source or "").lower()
    # Reject obvious project / compliance-sheet citations even if they mention a manual word later
    if re.search(
        r"\b(compliance\s+sheet|medical\s+centre|mobco|export\s+bank|el\s+sewedy|"
        r"project\s+specs?|consultant\s+comments?|full\s+comments)\b",
        t,
    ):
        return None
    if "programming" in t or "user guide" in t:
        return "programming"
    if "installation" in t or "instruction" in t:
        return "installation"
    if "catalog" in t:
        return "catalog"
    return None


def _page_in_manual(kind: str | None, page: str | None) -> bool:
    if not kind or not page:
        return False
    idx = ROOT / "data" / "knowledge" / "page_index.json"
    if not idx.exists():
        return False
    data = json.loads(idx.read_text(encoding="utf-8"))
    pages = data.get(kind) or []
    needle = str(page).strip()
    return any(str(p.get("printed_page")) == needle or str(p.get("pdf_page")) == needle for p in pages)


def _source_ignore_reason(source: str, kind: str | None, page: str | None) -> str | None:
    """
    Why schneider_sources was not updated.
    Past project sheets teach status/remarks only — never become the Source map.
    """
    raw = (source or "").strip()
    if not raw:
        return "no_source_text"
    if kind is None:
        return "not_schneider_manual"
    if not page:
        return "no_page_in_source"
    if not _page_in_manual(kind, page):
        return "page_not_in_manual_index"
    return None


def apply_corrections(
    rows: list[dict[str, str]],
    *,
    filename: str,
    offer_text: str = "",
) -> dict[str, Any]:
    """
    Client error loop: apply status + remarks to drive_capabilities, append to
    gold reference DB, and update schneider_sources only for verified manuals.
    """
    GOLD_DIR.mkdir(parents=True, exist_ok=True)
    patterns = load_patterns()
    family = resolve_pattern_key(offer_text) if offer_text else "ATV630"
    if family not in patterns:
        family = "ATV630"
    sources = load_schneider_sources()
    applied: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    ref_rows: list[dict[str, Any]] = []
    sources_updated = 0
    sources_ignored = 0
    ts = datetime.now(timezone.utc).isoformat()

    with GOLD_LOG.open("a", encoding="utf-8") as log:
        for row in rows:
            req = row.get("requirement") or ""
            cap = match_capability(req)
            status = _norm_status(row.get("status") or "")
            remark = shorten_remark(row.get("remarks") or "", 220)
            source = (row.get("source_document") or "").strip()
            record = {
                "ts": ts,
                "file": filename,
                "family": family,
                "capability_key": cap or "",
                "requirement": req[:400],
                "status": status,
                "remarks": remark,
                "source_document": source[:300],
            }
            log.write(json.dumps(record, ensure_ascii=False) + "\n")
            # Always feed gold DB so the same mistake is not repeated
            if req and status:
                ref_rows.append(
                    {
                        "requirement": req,
                        "status": status,
                        "remarks": remark,
                        "capability_key": cap or "",
                        "family": family,
                    }
                )
            if not cap:
                unmatched.append({"requirement": req[:180], "reason": "no capability match"})
                continue
            if status:
                patterns[family]["capabilities"].setdefault(cap, {})
                patterns[family]["capabilities"][cap]["status"] = status
            if remark:
                patterns[family]["capabilities"].setdefault(cap, {})
                patterns[family]["capabilities"][cap]["remark"] = remark
            page = _extract_page(source)
            kind = _doc_kind(source)
            ignore = _source_ignore_reason(source, kind, page)
            source_updated = False
            if ignore is None:
                meta = sources.get(cap) or {}
                doc_type = "Catalog"
                document = CATALOG_NAME
                if kind == "programming":
                    doc_type = "User guide"
                    document = "ATV600 Programming Manual (EAV64318)"
                elif kind == "installation":
                    doc_type = "Instruction sheet"
                    document = "ATV600 Installation Manual (EAV64301)"
                meta.update(
                    {
                        "doc_type": doc_type,
                        "document": document,
                        "page": page,
                        "evidence": f"Client correction in {filename}; page verified in {kind} index.",
                    }
                )
                sources[cap] = meta
                source_updated = True
                sources_updated += 1
            else:
                sources_ignored += 1
            applied.append(
                {
                    "capability_key": cap,
                    "status": status,
                    "remark_updated": bool(remark),
                    "source_updated": source_updated,
                    "source_ignored": ignore is not None,
                    "source_ignore_reason": ignore,
                }
            )

    PATTERNS_PATH.write_text(json.dumps(patterns, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    # Preserve _meta if present
    existing_meta = {}
    if SOURCES_PATH.exists():
        raw = json.loads(SOURCES_PATH.read_text(encoding="utf-8"))
        existing_meta = {k: v for k, v in raw.items() if str(k).startswith("_")}
    out_sources = {**existing_meta, **sources}
    SOURCES_PATH.write_text(json.dumps(out_sources, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    ref_appended = append_reference_rows(ref_rows, source_file=filename, family=family)

    return {
        "ok": True,
        "family": family,
        "parsed": len(rows),
        "applied": len(applied),
        "unmatched": len(unmatched),
        "sources_updated": sources_updated,
        "sources_ignored": sources_ignored,
        "reference_appended": ref_appended,
        "policy": (
            "Tab 3 is the client error loop: updates rule book + gold reference DB. "
            "Schneider Source map changes only for Catalog/Programming/Installation with verified pages. "
            "Past project sheets in data/references/ are never cited as Source."
        ),
        "applied_rows": applied,
        "unmatched_rows": unmatched[:40],
        "gold_log": str(GOLD_LOG.relative_to(ROOT)),
        "reference_db": str(REFERENCE_DB.relative_to(ROOT)),
    }
