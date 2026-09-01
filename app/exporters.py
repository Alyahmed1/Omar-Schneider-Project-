"""Excel and Word export — training-style compliance sheet with source Product Datasheet."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

# Excel/Word XML forbids most ASCII control chars (PDF snippets often include them).
_ILLEGAL_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _xml_safe(value: Any) -> str:
    if value is None:
        return ""
    return _ILLEGAL_XML.sub(" ", str(value))


def _datasheet_bits(attached_parts: list[dict[str, Any]] | None) -> list[str]:
    bits: list[str] = []
    for part in attached_parts or []:
        if not isinstance(part, dict):
            continue
        pn = _xml_safe(part.get("part_number", "")).strip()
        name = "Product Datasheet"
        for doc in part.get("documents") or []:
            if isinstance(doc, dict) and doc.get("file_name"):
                name = _xml_safe(doc.get("file_name"))
                break
        bits.append(f"{pn} → {name}" if pn else name)
    return bits


COLUMNS = [
    "Clause",
    "Requirement text",
    "Complies? (Yes/No/N/A)",
    "Why (short)",
    "Consultant page",
    "Schneider source (User guide / Instruction / Catalog)",
]


def _rows(draft: list[dict[str, Any]]) -> list[list[Any]]:
    rows = []
    for item in draft:
        rows.append(
            [
                _xml_safe(item.get("clause_id", "")),
                _xml_safe(item.get("requirement", "")),
                _xml_safe(item.get("status", "")),
                _xml_safe(item.get("remarks", "")),
                _xml_safe(item.get("source_page", "")),
                _xml_safe(item.get("source_document", "")),
            ]
        )
    return rows


def export_excel(
    draft: list[dict[str, Any]],
    family: str = "",
    attached_parts: list[dict[str, Any]] | None = None,
) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Compliance Sheet"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1F4E79")
    status_header_fill = PatternFill("solid", fgColor="0B6E4F")
    thin = Border(
        left=Side(style="thin", color="B0B0B0"),
        right=Side(style="thin", color="B0B0B0"),
        top=Side(style="thin", color="B0B0B0"),
        bottom=Side(style="thin", color="B0B0B0"),
    )
    wrap = Alignment(wrap_text=True, vertical="top")

    ws["A1"] = "Technical Submittal — Compliance Sheet"
    ws["A1"].font = Font(bold=True, size=14, color="1F4E79")
    ws.merge_cells("A1:F1")

    ws["A2"] = "Variable Frequency / Speed Drives (VFD / VSD)"
    ws["A2"].font = Font(bold=True, size=11)

    family = _xml_safe(family)
    ws["A3"] = f"Offered part numbers (from Part Lookup): {family}" if family else ""
    ws["A3"].font = Font(bold=True, size=11)

    attached_txt = ""
    bits = _datasheet_bits(attached_parts)
    if bits:
        attached_txt = "Schneider sources / Product Datasheets: " + "; ".join(bits)
    ws["A4"] = attached_txt
    ws.merge_cells("A4:F4")

    ws["A5"] = (
        "PRIMARY ANSWER — Complies?: Yes = complies; No = does not comply; N/A = not applicable. "
        "Why / page / source document support the answer."
    )
    ws["A5"].font = Font(bold=True, color="0B6E4F")
    ws.merge_cells("A5:F5")
    ws["A5"].alignment = wrap

    start = 7
    for col, name in enumerate(COLUMNS, start=1):
        cell = ws.cell(row=start, column=col, value=name)
        cell.font = header_font
        cell.fill = status_header_fill if col == 3 else header_fill
        cell.border = thin
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    yes_fill = PatternFill("solid", fgColor="C6EFCE")
    no_fill = PatternFill("solid", fgColor="FFC7CE")
    na_fill = PatternFill("solid", fgColor="FFEB9C")
    status_font = Font(bold=True, size=12)

    for r_idx, row in enumerate(_rows(draft), start=start + 1):
        for c_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.border = thin
            cell.alignment = wrap
            if c_idx == 3:
                cell.font = status_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                status = str(value).strip().upper()
                if status == "YES":
                    cell.fill = yes_fill
                elif status == "NO":
                    cell.fill = no_fill
                elif status in ("N/A", "NA"):
                    cell.fill = na_fill

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 48
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 40
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 42

    # Auto row height so wrapped Why / Source text does not overlay
    col_widths = {2: 48, 4: 40, 6: 42}
    for r_idx, row in enumerate(_rows(draft), start=start + 1):
        lines = 2
        for c_idx, width in col_widths.items():
            text = str(row[c_idx - 1] or "")
            chunk_lines = 0
            for paragraph in text.split("\n") or [""]:
                chunk_lines += max(1, (len(paragraph) + width - 1) // width)
            lines = max(lines, chunk_lines)
        ws.row_dimensions[r_idx].height = min(max(18 + lines * 14, 36), 180)

    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def export_word(
    draft: list[dict[str, Any]],
    family: str = "",
    attached_parts: list[dict[str, Any]] | None = None,
) -> bytes:
    """
    Training-sheet style Word export:
    header block + requirement paragraphs with Comply answer, why, page, source datasheet.
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # Header like training PDF
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Technical Submittal")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Compliance Sheet")
    r2.bold = True
    r2.font.size = Pt(14)

    sub = doc.add_paragraph()
    sr = sub.add_run("Variable Frequency / Speed Drives (VFD / VSD)")
    sr.bold = True
    sr.font.size = Pt(11)

    family = _xml_safe(family)
    if family:
        op = doc.add_paragraph()
        orun = op.add_run(f"Offered part numbers (from Part Lookup): {family}")
        orun.bold = True

    bits = _datasheet_bits(attached_parts)
    if bits:
        doc.add_paragraph("Schneider sources / Product Datasheets: " + "; ".join(bits))

    legend = doc.add_paragraph()
    lr = legend.add_run(
        "PRIMARY ANSWER — Complies?: Yes = complies; No = does not comply; N/A = not applicable."
    )
    lr.bold = True
    lr.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)

    doc.add_paragraph("")

    # Training-like body: each clause as a block (requirement + Comply + why + sources)
    for item in draft:
        req = _xml_safe(item.get("requirement", "")).strip()
        status = _xml_safe(item.get("status", "Yes")).strip()
        remarks = _xml_safe(item.get("remarks", "")).strip()
        page = _xml_safe(item.get("source_page", "")).strip()
        source_doc = _xml_safe(item.get("source_document", "")).strip()
        clause_id = _xml_safe(item.get("clause_id", "")).strip()

        block = doc.add_paragraph()
        if clause_id:
            cr = block.add_run(f"{clause_id}  ")
            cr.bold = True
            cr.font.size = Pt(10)
        rr = block.add_run(req)
        rr.font.size = Pt(10)

        ans = doc.add_paragraph()
        # Match training tone: Comply / Yes with justification
        if status.upper() == "YES":
            label = "Comply (Yes)"
            color = RGBColor(0x0B, 0x6E, 0x4F)
        elif status.upper() == "NO":
            label = "Does not comply (No)"
            color = RGBColor(0xA0, 0x2A, 0x2A)
        else:
            label = "Not applicable (N/A)"
            color = RGBColor(0x8A, 0x6D, 0x00)
        ar = ans.add_run(label)
        ar.bold = True
        ar.font.size = Pt(11)
        ar.font.color.rgb = color
        if remarks:
            ans.add_run(" — ").font.size = Pt(10)
            wr = ans.add_run(remarks)
            wr.font.size = Pt(10)

        # Resource / source block — readable size (not gray italic footnote)
        if page or source_doc:
            meta = doc.add_paragraph()
            meta.paragraph_format.space_before = Pt(4)
            meta.paragraph_format.space_after = Pt(2)

            def _meta_label(text: str) -> None:
                run = meta.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

            def _meta_value(text: str) -> None:
                run = meta.add_run(text)
                run.bold = False
                run.italic = False
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

            if page:
                _meta_label("Consultant page: ")
                _meta_value(page)
            if source_doc:
                if page:
                    meta.add_run().add_break()
                _meta_label("Schneider source: ")
                _meta_value(source_doc)

        # light spacer
        doc.add_paragraph("")

    note = doc.add_paragraph()
    note.add_run(
        "Draft generated for engineering review. Verify all Yes/No/N/A answers before formal submittal."
    ).italic = True

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def save_exports(
    draft: list[dict[str, Any]],
    family: str,
    out_dir: str | Path,
    stem: str = "compliance_draft",
    attached_parts: list[dict[str, Any]] | None = None,
) -> dict[str, str]:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    xlsx_path = out_dir / f"{stem}.xlsx"
    docx_path = out_dir / f"{stem}.docx"
    xlsx_path.write_bytes(export_excel(draft, family, attached_parts))
    docx_path.write_bytes(export_word(draft, family, attached_parts))
    return {"excel": str(xlsx_path), "word": str(docx_path)}
